import os
import re
from typing import Optional

# Import libraries for document parsing
import pdfplumber
from docx import Document

def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Extract text from a document file (PDF, DOCX, or TXT)
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text or None if extraction failed
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == ".pdf":
            return extract_text_from_pdf(file_path)
        elif file_extension == ".docx":
            return extract_text_from_docx(file_path)
        elif file_extension == ".txt":
            return extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return None

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using pdfplumber
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text
    """
    text = ""
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n\n"
    
    # Clean up text
    text = clean_text(text)
    
    return text

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text
    """
    doc = Document(file_path)
    
    # Extract text from paragraphs
    paragraphs = [para.text for para in doc.paragraphs]
    
    # Extract text from tables
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            tables_text.append(" | ".join(row_text))
    
    # Combine all text
    all_text = "\n\n".join(paragraphs) + "\n\n" + "\n".join(tables_text)
    
    # Clean up text
    text = clean_text(all_text)
    
    return text

def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a plain text file
    
    Args:
        file_path: Path to the text file
        
    Returns:
        Extracted text
    """
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    
    # Clean up text
    text = clean_text(text)
    
    return text

def clean_text(text: str) -> str:
    """
    Clean up extracted text
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Replace multiple spaces with a single space
    text = re.sub(r' {2,}', ' ', text)
    
    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n\t]', '', text)
    
    return text.strip()